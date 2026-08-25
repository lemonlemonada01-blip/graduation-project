import os
import re
from pathlib import Path
from typing import List, Dict, Any, Union, Optional
import pypdf
import docx

CODE_EXTENSIONS = {
    '.py', '.cpp', '.c', '.h', '.hpp', '.cc', '.cxx',
    '.java', '.js', '.ts', '.jsx', '.tsx', '.cs',
    '.go', '.rs', '.php', '.html', '.css', '.scss',
    '.rb', '.swift', '.kt', '.kts', '.sql', '.sh',
    '.bash', '.zsh', '.pl', '.r', '.m', '.scala',
    '.dart', '.vue', '.svelte', '.lua', '.yaml', '.yml',
    '.json', '.xml', '.toml'
}

TEXT_EXTENSIONS = {
    '.txt', '.pdf', '.docx', '.doc', '.md', '.markdown',
    '.rst', '.tex', '.rtf', '.log', '.csv', '.tsv'
}

BINARY_EXTENSIONS = {
    '.png', '.jpg', '.jpeg', '.gif', '.ico', '.bmp', '.svg',
    '.exe', '.dll', '.so', '.dylib', '.db', '.sqlite', '.sqlite3',
    '.pyc', '.class', '.jar', '.war', '.zip', '.tar', '.gz', '.7z',
    '.rar', '.pdf.tmp', '.bin', '.dat', '.o', '.a', '.obj', '.eot',
    '.ttf', '.woff', '.woff2'
}

DEFAULT_IGNORE_DIRS = {
    # Python
    'venv', 'env', '__pycache__', 'htmlcov', 'wheels',
    # Node.js / Web
    'node_modules', 'bower_components', 'jspm_packages',
    # Java/Android/Kotlin/Scala
    'captures', 'gen',
    # C/C++/.NET
    'Debug', 'Release', 'x64', 'x86', 'cmake-build-debug', 'cmake-build-release', 'CMakeFiles',
    # Go / Rust / Ruby / PHP
    'vendor', 'pkg', 'tmp',
    # Mobile (Flutter / iOS)
    'Pods', 'ephemeral', 'DerivedData',
    # Unity / GameDev
    'Library', 'Temp', 'Obj', 'Builds', 'Logs', 'MemoryCaptures',
    # General Build/Outputs/Cache
    'build', 'dist', 'target', 'bin', 'obj', 'out', 'coverage',
    # Project Specific
    'all-MiniLM-L6-v2', 'models', 'chroma_db', 'chroma_demo_db', 'chroma_test_folder_db', 'chroma_app_db'
}

# Auto-generated configs, lockfiles, and generic repository documentation files
DEFAULT_IGNORE_FILES = {
    # Locks
    'package-lock.json', 'yarn.lock', 'pnpm-lock.yaml', 'bun.lockb',
    'pipfile.lock', 'poetry.lock', 'gemfile.lock', 'composer.lock', 
    'cargo.lock', 'go.sum', 'pubspec.lock', 'podfile.lock',
    # Configs
    'config.json', '.gitignore', 'tsconfig.json', 'jsconfig.json', 
    '.eslintrc', '.prettierrc', 'babel.config.js', 'webpack.config.js', 
    'vite.config.js', 'cmakelists.txt', 'makefile', 'setup.cfg', 'manifest.in',
    '.ds_store', 'thumbs.db', 'dockerfile', 'docker-compose.yml',
    # Generic documentation & repository meta files
    'readme.md', 'readme.txt', 'readme.rst', 'readme.markdown', 'readme',
    'license', 'license.txt', 'license.md', 'copying',
    'changelog.md', 'contributing.md', 'code_of_conduct.md', 'security.md',
    # Flutter & Mobile auto-generated files
    'generated_plugin_registrant.dart', 'generated_plugin_registrant.h',
    'generated_plugin_registrant.m', 'generatedpluginregistrant.swift'
}

def sanitize_text(text: str) -> str:
    """Removes NUL (0x00) characters and invalid unicode bytes for PostgreSQL compatibility."""
    if not text:
        return ""
    return text.replace('\x00', '')

class FileExtractor:
    """Extracts text/code from files and project directories with encoding fallbacks and directory/config filtering."""

    @staticmethod
    def extract_text_from_pdf(pdf_path: Union[str, Path]) -> str:
        """Extract text from a PDF file using pypdf."""
        pdf_path = Path(pdf_path)
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        text_chunks = []
        with open(pdf_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
        return sanitize_text("\n".join(text_chunks))

    @staticmethod
    def extract_text_from_docx(docx_path: Union[str, Path]) -> str:
        """Extract text from a Word document (.docx) using python-docx."""
        docx_path = Path(docx_path)
        if not docx_path.exists():
            raise FileNotFoundError(f"Word document not found: {docx_path}")

        doc = docx.Document(docx_path)
        text_chunks = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_chunks.append(paragraph.text)
                
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_chunks.append(" | ".join(row_text))
                    
        return sanitize_text("\n".join(text_chunks))

    @staticmethod
    def read_text_file(file_path: Union[str, Path]) -> str:
        """Read text/code file using multiple encoding fallbacks (UTF-8, UTF-8-sig, CP1256, Latin-1)."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        encodings = ['utf-8', 'utf-8-sig', 'cp1256', 'latin-1', 'iso-8859-1']
        raw_text = ""
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    raw_text = f.read()
                    break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if not raw_text:
            try:
                with open(file_path, 'rb') as f:
                    raw_text = f.read().decode('utf-8', errors='ignore')
            except Exception:
                raw_text = ""

        return sanitize_text(raw_text)

    @classmethod
    def categorize_file(cls, filename_or_path: Union[str, Path]) -> str:
        """Categorize file as 'code' or 'text' based on file extension."""
        ext = Path(filename_or_path).suffix.lower()
        if ext in CODE_EXTENSIONS:
            return 'code'
        elif ext in TEXT_EXTENSIONS:
            return 'text'
        return 'text'

    @classmethod
    def extract_file(cls, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract content and metadata from a single file."""
        path = Path(file_path)
        ext = path.suffix.lower()
        file_type = cls.categorize_file(path)

        if ext == '.pdf':
            content = cls.extract_text_from_pdf(path)
        elif ext in ('.docx', '.doc'):
            content = cls.extract_text_from_docx(path)
        else:
            content = cls.read_text_file(path)

        return {
            "path": str(path),
            "filename": path.name,
            "extension": ext,
            "file_type": file_type,
            "content": sanitize_text(content)
        }

    @classmethod
    def scan_project_directory(
        cls, 
        dir_path: Union[str, Path], 
        ignore_dirs: Optional[set] = None,
        ignore_files: Optional[set] = None,
        min_lines: int = 3
    ) -> List[Dict[str, Any]]:
        """
        Recursively scan a project directory, filtering out dependency, build, binary, documentation, and trivial init files.
        """
        dir_path = Path(dir_path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Directory not found: {dir_path}")

        ignored_d = DEFAULT_IGNORE_DIRS if ignore_dirs is None else set(ignore_dirs)
        ignored_f = DEFAULT_IGNORE_FILES if ignore_files is None else set(ignore_files)
        extracted_files = []
        
        total_size = 0
        MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
        MAX_TOTAL_SIZE = 500 * 1024 * 1024 # 500MB

        for root, dirs, files in os.walk(dir_path):
            dirs[:] = [d for d in dirs if d not in ignored_d and not d.startswith('.')]
            
            for file in files:
                file_name_lower = file.lower()
                
                # Filter out boilerplate, documentation, config files, and generated dart files
                if file_name_lower in ignored_f or file_name_lower.startswith('config.'):
                    continue
                
                ignored_suffixes = (
                    '.g.dart', '.freezed.dart', '.min.js', '.min.css', '.map', 
                    '.pbxproj', '.xcodeproj', '.xcworkspace', '.suo', '.user', 
                    '.sln', '.csproj', '.vcxproj'
                )
                if file_name_lower.endswith(ignored_suffixes):
                    continue

                file_path = Path(root) / file
                ext = file_path.suffix.lower()

                # Ignore known binary file extensions
                if ext in BINARY_EXTENSIONS:
                    continue

                if ext in CODE_EXTENSIONS or ext in TEXT_EXTENSIONS:
                    try:
                        file_size = file_path.stat().st_size
                        if file_size > MAX_FILE_SIZE:
                            continue
                        if total_size + file_size > MAX_TOTAL_SIZE:
                            break
                        
                        record = cls.extract_file(file_path)
                        lines = [l.strip() for l in record['content'].splitlines() if l.strip() and not l.strip().startswith('#')]
                        
                        # Filter out trivial files with fewer than min_lines
                        if len(lines) < min_lines:
                            continue

                        # Filter out trivial __init__.py files with < 5 non-trivial lines
                        if file_name_lower == '__init__.py' and len(lines) < 5:
                            continue

                        record["relative_path"] = str(file_path.relative_to(dir_path))
                        extracted_files.append(record)
                        total_size += file_size
                    except Exception:
                        continue

        return extracted_files
